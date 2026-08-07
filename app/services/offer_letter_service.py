"""Offer Letter Generator Service.

Generates professional offer letters using:
- Jinja2 HTML templates → PDF (via weasyprint)
- python-docx for DOCX export
- AI-enhanced content via Ollama

Supports:
- Editable Jinja2 templates
- PDF export
- DOCX export
- Variable substitution
- Custom company branding
"""

from __future__ import annotations

import logging
import os
import uuid
from datetime import datetime, timezone
from typing import Any

from app.core.config import settings
from app.llm.client import LLMClient, get_llm_client
from app.llm.prompts import PromptLibrary
from app.llm.response_parser import ResponseParser

logger = logging.getLogger(__name__)

os.makedirs(settings.OFFER_LETTER_DIR, exist_ok=True)


class OfferLetterContext:
    """Data context for offer letter generation."""

    def __init__(
        self,
        candidate_name: str,
        position: str,
        department: str,
        salary: str,
        joining_date: str,
        company_name: str,
        reporting_to: str = "Hiring Manager",
        location: str = "Head Office",
        benefits: list[str] | None = None,
        employment_type: str = "Full-Time",
        probation_months: int = 3,
        offer_expiry_days: int = 7,
        additional_terms: str = "",
        hr_signatory: str = "HR Department",
        custom_fields: dict[str, str] | None = None,
    ) -> None:
        self.candidate_name = candidate_name
        self.position = position
        self.department = department
        self.salary = salary
        self.joining_date = joining_date
        self.company_name = company_name
        self.reporting_to = reporting_to
        self.location = location
        self.benefits = benefits or [
            "Competitive salary package",
            "Health and dental insurance",
            "Flexible working hours",
            "Professional development budget",
        ]
        self.employment_type = employment_type
        self.probation_months = probation_months
        self.offer_expiry_days = offer_expiry_days
        self.additional_terms = additional_terms
        self.hr_signatory = hr_signatory
        self.custom_fields = custom_fields or {}
        self.generated_at = datetime.now(tz=timezone.utc).strftime("%B %d, %Y")


class OfferLetterService:
    """Service for generating PDF and DOCX offer letters."""

    def __init__(self, llm_client: LLMClient | None = None) -> None:
        self._llm = llm_client or get_llm_client()

    async def generate_offer_letter(
        self,
        context: OfferLetterContext,
        use_ai_content: bool = True,
        model: str | None = None,
    ) -> dict[str, str]:
        """Generate offer letter text content.

        Args:
            context: Offer letter data context.
            use_ai_content: Whether to use AI to enhance content.

        Returns:
            Dict with 'subject', 'full_letter_text', 'key_terms'.
        """
        if use_ai_content:
            return await self._generate_with_ai(context, model)
        else:
            return self._generate_from_template(context)

    async def export_pdf(
        self,
        content: str,
        context: OfferLetterContext,
    ) -> str:
        """Export offer letter content to PDF.

        Returns: File path to generated PDF.
        """
        file_name = f"offer_letter_{uuid.uuid4().hex[:8]}.pdf"
        file_path = os.path.join(settings.OFFER_LETTER_DIR, file_name)

        html_content = self._build_html(content, context)

        try:
            import weasyprint
            weasyprint.HTML(string=html_content).write_pdf(file_path)
            logger.info("Offer letter PDF generated: %s", file_path)
        except ImportError:
            logger.warning("weasyprint not installed — writing plain text instead")
            txt_path = file_path.replace(".pdf", ".txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(content)
            return txt_path
        except Exception as exc:
            logger.error("PDF generation failed: %s", exc)
            # Fallback: save HTML
            html_path = file_path.replace(".pdf", ".html")
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            return html_path

        return file_path

    def export_docx(
        self,
        content: str,
        context: OfferLetterContext,
    ) -> str:
        """Export offer letter content to DOCX.

        Returns: File path to generated DOCX.
        """
        file_name = f"offer_letter_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(settings.OFFER_LETTER_DIR, file_name)

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Header
            title = doc.add_heading(f"OFFER LETTER — {context.company_name}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Date
            date_para = doc.add_paragraph(f"Date: {context.generated_at}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_paragraph()

            # Candidate address
            doc.add_paragraph(f"Dear {context.candidate_name},")
            doc.add_paragraph()

            # Main content paragraphs
            for paragraph in content.split("\n\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph(context.hr_signatory)
            doc.add_paragraph(context.company_name)

            doc.save(file_path)
            logger.info("Offer letter DOCX generated: %s", file_path)

        except ImportError:
            logger.warning("python-docx not installed — writing plain text instead")
            txt_path = file_path.replace(".docx", ".txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(content)
            return txt_path
        except Exception as exc:
            logger.error("DOCX generation failed: %s", exc)

        return file_path

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    async def _generate_with_ai(
        self,
        context: OfferLetterContext,
        model: str | None = None,
    ) -> dict[str, str]:
        """Use LLM to generate professional offer letter content."""
        response = await self._llm.complete(
            prompt=PromptLibrary.offer_letter_user(
                candidate_name=context.candidate_name,
                position=context.position,
                department=context.department,
                salary=context.salary,
                joining_date=context.joining_date,
                company_name=context.company_name,
                benefits=context.benefits,
                reporting_to=context.reporting_to,
            ),
            system=PromptLibrary.OFFER_LETTER_SYSTEM,
            model=model,
            json_mode=True,
            temperature=0.3,
            num_predict=2000,
        )

        data = ResponseParser.extract_json_object(response) if response else {}
        if data:
            return {
                "subject": ResponseParser.get_str(data, "subject",
                                                   f"Offer of Employment — {context.position}"),
                "full_letter_text": ResponseParser.get_str(data, "full_letter_text",
                                                            self._generate_from_template(context)["full_letter_text"]),
                "key_terms": data.get("key_terms", {}),
            }

        # Fallback to template
        return self._generate_from_template(context)

    @staticmethod
    def _generate_from_template(context: OfferLetterContext) -> dict[str, str]:
        """Generate offer letter from built-in template (no LLM)."""
        benefits_text = "\n".join(f"  • {b}" for b in context.benefits)

        letter_text = f"""Dear {context.candidate_name},

We are delighted to offer you the position of {context.position} in the {context.department} department at {context.company_name}. After a thorough review of your experience and qualifications, we are confident that you will be a valuable addition to our team.

POSITION DETAILS
----------------
Position: {context.position}
Department: {context.department}
Reporting To: {context.reporting_to}
Location: {context.location}
Employment Type: {context.employment_type}
Start Date: {context.joining_date}

COMPENSATION
------------
Annual Salary: {context.salary}
Probation Period: {context.probation_months} months

BENEFITS
--------
{benefits_text}

TERMS
-----
This offer is contingent upon the successful completion of background verification and reference checks. This offer letter is valid for {context.offer_expiry_days} days from the date of issue.

{context.additional_terms}

We look forward to welcoming you to the {context.company_name} team. Please sign and return a copy of this letter to confirm your acceptance.

Warm regards,
{context.hr_signatory}
{context.company_name}
Date: {context.generated_at}"""

        return {
            "subject": f"Offer of Employment — {context.position} at {context.company_name}",
            "full_letter_text": letter_text,
            "key_terms": {
                "position": context.position,
                "department": context.department,
                "start_date": context.joining_date,
                "salary": context.salary,
                "reporting_to": context.reporting_to,
            },
        }

    @staticmethod
    def _build_html(content: str, context: OfferLetterContext) -> str:
        """Wrap offer letter content in professional HTML template."""
        paragraphs = "".join(
            f"<p>{p.strip()}</p>"
            for p in content.split("\n\n")
            if p.strip()
        )
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
  body {{ font-family: 'Georgia', serif; font-size: 11pt; color: #222; margin: 60px; line-height: 1.6; }}
  h1 {{ color: #0f172a; text-align: center; font-size: 16pt; }}
  .header {{ text-align: center; border-bottom: 2px solid #0f172a; padding-bottom: 20px; margin-bottom: 30px; }}
  .date {{ text-align: right; color: #555; margin-bottom: 20px; }}
  p {{ margin: 12px 0; }}
  .footer {{ margin-top: 40px; border-top: 1px solid #ccc; padding-top: 20px; }}
</style>
</head>
<body>
  <div class="header">
    <h1>{context.company_name}</h1>
    <p>Official Offer of Employment</p>
  </div>
  <div class="date">Date: {context.generated_at}</div>
  <div class="content">
    {paragraphs}
  </div>
  <div class="footer">
    <p>{context.hr_signatory}</p>
    <p>{context.company_name}</p>
  </div>
</body>
</html>"""


# ---------------------------------------------------------------------------
# Singleton
# ---------------------------------------------------------------------------

_offer_service: OfferLetterService | None = None


def get_offer_letter_service() -> OfferLetterService:
    """Return the global offer letter service singleton."""
    global _offer_service
    if _offer_service is None:
        _offer_service = OfferLetterService()
    return _offer_service
