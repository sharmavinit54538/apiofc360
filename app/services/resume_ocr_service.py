"""Resume OCR Service for extracting raw text from PDF, DOCX, DOC, TXT, PNG, JPG, JPEG, and TIFF files."""

from __future__ import annotations

import io
import logging
import os
import re
from typing import Any

from app.core.exceptions import AppException
from app.services.google_document_ai_service import GoogleDocumentAIService

logger = logging.getLogger(__name__)


class ResumeOCRService:
    """Service for extracting text from resume files using Google Document AI, pypdf, python-docx, and OCR fallback."""

    def __init__(self, doc_ai_service: GoogleDocumentAIService | None = None) -> None:
        self.doc_ai_service = doc_ai_service or GoogleDocumentAIService()

    async def extract_text(self, file_bytes: bytes, file_name: str, mime_type: str) -> dict[str, Any]:
        """Extract full raw text from uploaded resume using Document AI with robust format fallbacks."""
        ext = os.path.splitext(file_name)[1].lower()
        logger.info("Extracting resume text | file_name=%s | ext=%s | mime=%s", file_name, ext, mime_type)

        # 1. Try Google Document AI if configured
        try:
            doc_ai_result = await self.doc_ai_service.process_document(file_bytes, mime_type)
            document_obj = doc_ai_result.get("document")
            extracted_text = getattr(document_obj, "text", "") or ""
            if extracted_text and len(extracted_text.strip()) > 30:
                logger.info("Document AI OCR successfully extracted %s chars", len(extracted_text))
                return {
                    "raw_text": extracted_text.strip(),
                    "ocr_engine": "google_document_ai",
                    "confidence": 0.98,
                }
        except Exception as exc:
            logger.debug("Google Document AI processing skipped or unconfigured: %s", exc)

        # 2. Format-Specific Fallback Extraction
        raw_text = ""
        engine_used = "format_fallback"
        confidence = 0.85

        if ext == ".pdf" or mime_type == "application/pdf":
            raw_text, engine_used, confidence = self._extract_pdf(file_bytes)
        elif ext in [".docx", ".doc"] or "wordprocessingml" in mime_type or "msword" in mime_type:
            raw_text = self._extract_docx(file_bytes)
            engine_used = "docx_fallback"
            confidence = 0.92
        elif ext == ".txt" or "text/plain" in mime_type:
            raw_text = self._extract_txt(file_bytes)
            engine_used = "txt_decoder"
            confidence = 0.95
        elif ext in [".png", ".jpg", ".jpeg", ".tif", ".tiff"] or "image" in mime_type:
            raw_text = self._extract_image_fallback(file_bytes, ext=ext)
            engine_used = "image_ocr_engine"
            confidence = 0.88

        if not raw_text or len(raw_text.strip()) < 10:
            # Final fallback text decoding with safe encodings
            raw_text = self._extract_txt(file_bytes)
            if raw_text and len(raw_text.strip()) >= 10:
                engine_used = "raw_text_decoder"

        if not raw_text or len(raw_text.strip()) == 0:
            raise AppException(
                message="Unable to extract text from the uploaded resume file. The file may be corrupt, unreadable, or password-protected.",
                status_code=400,
            )

        logger.info("Resume text extracted %s chars using %s", len(raw_text), engine_used)
        return {
            "raw_text": raw_text.strip(),
            "ocr_engine": engine_used,
            "confidence": confidence,
        }

    def _extract_pdf(self, file_bytes: bytes) -> tuple[str, str, float]:
        """Extract text from PDF using pypdf. If scanned (insufficient text), fall back to OCR on page images."""
        text_pages: list[str] = []
        has_images = False

        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page_idx, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_pages.append(page_text.strip())
                else:
                    # Page has no selectable text — check for embedded images
                    if hasattr(page, "images") and page.images:
                        has_images = True
                        for img in page.images:
                            img_text = self._extract_image_fallback(img.data, ext=".png")
                            if img_text.strip():
                                text_pages.append(img_text.strip())

            combined_text = "\n\n".join(text_pages).strip()
            if len(combined_text) >= 50:
                return combined_text, "pypdf_native", 0.95

            # If combined selectable text is too short (< 50 chars) and we haven't checked images yet:
            if not has_images and len(reader.pages) > 0:
                ocr_pages = []
                for page in reader.pages:
                    if hasattr(page, "images"):
                        for img in page.images:
                            img_text = self._extract_image_fallback(img.data, ext=".png")
                            if img_text.strip():
                                ocr_pages.append(img_text.strip())
                if ocr_pages:
                    return "\n\n".join(ocr_pages).strip(), "pypdf_ocr_fallback", 0.85

            if combined_text:
                return combined_text, "pypdf_sparse", 0.70

        except Exception as exc:
            logger.warning("pypdf extraction exception: %s. Trying regex text recovery.", exc)

        # Binary string extraction fallback
        text_content = file_bytes.decode("latin-1", errors="ignore")
        readable = re.findall(r"[A-Za-z0-9\s.,@+\-(){}:;/]{4,}", text_content)
        extracted = "\n".join(readable).strip()
        return extracted, "pdf_binary_recovery", 0.50

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Extract paragraphs, headings, and tables from DOCX document."""
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            content_blocks: list[str] = []

            # 1. Extract paragraphs with heading markers
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading" in style_name or "title" in style_name:
                    content_blocks.append(f"## {text}")
                else:
                    content_blocks.append(text)

            # 2. Extract tables row by row with delimiters
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # Deduplicate adjacent duplicate merged cells
                    unique_cells = []
                    for c in cells:
                        if not unique_cells or unique_cells[-1] != c:
                            unique_cells.append(c)
                    if unique_cells:
                        content_blocks.append(" | ".join(unique_cells))

            if content_blocks:
                return "\n".join(content_blocks)

        except Exception as exc:
            logger.debug("python-docx extraction failed: %s. Trying zip XML fallback.", exc)

        # Fallback zip file XML reading
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                texts = [node.text for node in tree.iter() if node.text]
                return "\n".join(texts)
        except Exception as exc:
            logger.warning("DOCX zip XML extraction failed: %s", exc)

        # Binary regex fallback for legacy .doc
        text_content = file_bytes.decode("latin-1", errors="ignore")
        readable = re.findall(r"[A-Za-z0-9\s.,@+\-(){}:;/]{4,}", text_content)
        return "\n".join(readable).strip()

    def _extract_txt(self, file_bytes: bytes) -> str:
        """Extract plain text from bytes using safe encoding detection/fallbacks."""
        encodings = ["utf-8", "utf-8-sig", "utf-16", "latin-1", "cp1252", "ascii"]
        for enc in encodings:
            try:
                decoded = file_bytes.decode(enc)
                if decoded and len(decoded.strip()) > 0:
                    return decoded.strip()
            except (UnicodeDecodeError, UnicodeError):
                continue
        return file_bytes.decode("utf-8", errors="ignore").strip()

    def _extract_image_fallback(self, file_bytes: bytes, ext: str = ".png") -> str:
        """Run OCR on image bytes using available OCR engines."""
        # 1. Try OCR Engine Selector
        try:
            from app.ocr.engine_selector import get_ocr_selector
            selector = get_ocr_selector()
            res = selector.extract_text_from_image_bytes(file_bytes, ext=ext)
            if res and res.success and len(res.text.strip()) > 10:
                return res.text.strip()
        except Exception as exc:
            logger.debug("OCREngineSelector failed: %s", exc)

        # 2. Try pytesseract / PIL directly if installed
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(io.BytesIO(file_bytes))
            return pytesseract.image_to_string(img).strip()
        except Exception:
            pass

        return ""
