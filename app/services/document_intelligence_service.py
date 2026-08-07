"""Document Intelligence Orchestrator Service.

Coordinates document pre-processing, multi-engine OCR, AI classification,
key-value JSON extraction, validation rules, risk/compliance analysis,
and version comparisons.
"""

from __future__ import annotations

import hashlib
import logging
import os
import uuid
from typing import Any, Optional

from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.llm.client import get_llm_client
from app.llm.prompts import PromptLibrary
from app.llm.response_parser import ResponseParser
from app.ocr.doc_ocr_orchestrator import get_ocr_orchestrator
from app.validators.doc_validator import DocumentValidator

logger = logging.getLogger(__name__)


# Standard schemas for typical business documents to feed into Ollama
SCHEMAS = {
    "PAN_CARD": """{
  "name": "Full Name as printed",
  "father_name": "Father's Name as printed",
  "dob": "Date of Birth (YYYY-MM-DD)",
  "pan_number": "10-character PAN number",
  "issue_date": "Issue date if printed else null"
}""",
    "AADHAAR": """{
  "name": "Full Name",
  "dob": "DOB or Year of Birth",
  "gender": "MALE/FEMALE/OTHER",
  "aadhaar_number": "12-digit Aadhaar number with format XXXX XXXX XXXX",
  "address": "Address if printed else null"
}""",
    "GST_DOCUMENT": """{
  "legal_name": "Legal name of business",
  "trade_name": "Trade name or DBA",
  "gstin": "15-character GSTIN",
  "constitution": "Proprietorship/Partnership/Company",
  "registration_date": "Date of registration",
  "address": "Principal place of business address"
}""",
    "INVOICE": """{
  "invoice_number": "Invoice number",
  "invoice_date": "Invoice issue date",
  "vendor_name": "Company/Vendor name",
  "vendor_gstin": "Vendor GSTIN if present",
  "buyer_name": "Client/Buyer name",
  "buyer_gstin": "Buyer GSTIN if present",
  "items": [
    {
      "description": "Item description",
      "quantity": 1.0,
      "rate": 100.0,
      "amount": 100.0
    }
  ],
  "subtotal": 0.0,
  "tax_amount": 0.0,
  "total_amount": 0.0,
  "currency": "INR/USD/etc"
}""",
    "PASSPORT": """{
  "passport_number": "Passport number",
  "given_name": "Given names",
  "surname": "Surname",
  "nationality": "Nationality",
  "dob": "DOB (YYYY-MM-DD)",
  "place_of_birth": "Birth city/state",
  "place_of_issue": "Issue office/city",
  "issue_date": "Issue date (YYYY-MM-DD)",
  "expiry_date": "Expiry date (YYYY-MM-DD)"
}""",
    "CONTRACT": """{
  "parties": ["Party A Legal Name", "Party B Legal Name"],
  "agreement_date": "Effective date of agreement",
  "governing_law": "Jurisdiction/State/Country",
  "termination_clause": "Summary of termination rights",
  "payment_terms": "Payment deadlines or cycles"
}""",
    "RESUME": """{
  "name": "Candidate Full Name",
  "email": "Email",
  "phone": "Phone number",
  "skills": ["Python", "FastAPI"],
  "experience_years": 5.0,
  "education": ["Degree from University"]
}"""
}


class DocumentIntelligenceService:
    """Orchestrates the lifecycle of Document processing and analysis."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.llm = get_llm_client()
        self.ocr = get_ocr_orchestrator()

    async def register_document(
        self,
        file_path: str,
        file_name: str,
        uploaded_by: Optional[uuid.UUID] = None,
        company_id: Optional[uuid.UUID] = None,
    ) -> tuple[str, bool]:
        """Verify checksum, check for duplicates, and insert AnalyzedDocument.

        Returns (doc_id, is_duplicate).
        """
        # Compute SHA-256 checksum
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        checksum = hasher.hexdigest()

        # Check duplicate
        from app.models.ai_document_analysis import AnalyzedDocument
        stmt = select(AnalyzedDocument).where(
            AnalyzedDocument.file_checksum == checksum,
            AnalyzedDocument.status != "FAILED"
        )
        res = await self.db.execute(stmt)
        existing = res.scalar_one_or_none()

        if existing:
            # Audit log duplicate discovery
            await self.log_audit(
                user_id=uploaded_by,
                doc_id=existing.id,
                action="UPLOAD",
                details=f"Duplicate document uploaded: {file_name}. Matches existing ID: {existing.id}"
            )
            return str(existing.id), True

        # Insert new record
        _, ext = os.path.splitext(file_name.lower())
        new_doc = AnalyzedDocument(
            id=uuid.uuid4(),
            company_id=company_id,
            file_path=file_path,
            file_name=file_name,
            file_size=os.path.getsize(file_path),
            file_type=ext.lstrip("."),
            file_checksum=checksum,
            status="PENDING",
            uploaded_by=uploaded_by,
        )
        self.db.add(new_doc)
        await self.db.commit()

        # Save initial version
        from app.models.ai_document_analysis import DocumentAnalysisVersion
        version = DocumentAnalysisVersion(
            document_id=new_doc.id,
            version_number=1,
            file_path=file_path,
            file_name=file_name,
            file_size=new_doc.file_size,
            file_checksum=checksum,
            uploaded_by=uploaded_by,
        )
        self.db.add(version)
        await self.db.commit()

        await self.log_audit(
            user_id=uploaded_by,
            doc_id=new_doc.id,
            action="UPLOAD",
            details=f"Registered document {file_name} with version 1"
        )

        return str(new_doc.id), False

    async def process_document(
        self,
        doc_uuid: uuid.UUID,
        model: Optional[str] = None,
    ) -> dict[str, Any]:
        """Orchestrate OCR, Classification, and key-value JSON Extraction."""
        from app.models.ai_document_analysis import AnalyzedDocument

        stmt = select(AnalyzedDocument).where(AnalyzedDocument.id == doc_uuid)
        res = await self.db.execute(stmt)
        doc = res.scalar_one_or_none()
        if not doc:
            raise ValueError(f"Document with ID {doc_uuid} not found")

        # 1. Run OCR (if PDF/images need it)
        doc.status = "PROCESSING"
        await self.db.commit()

        ocr_res = self.ocr.extract_text(doc.file_path)
        if not ocr_res.is_success and not doc.raw_text:
            # If text parsing from original document exists (like standard PDF) we can skip OCR failures
            if doc.file_type.lower() == "pdf":
                doc.raw_text = self._extract_pdf_text_fallback(doc.file_path)
            elif doc.file_type.lower() == "docx":
                doc.raw_text = self._extract_docx_text_fallback(doc.file_path)
            elif doc.file_type.lower() == "txt":
                with open(doc.file_path, "r", encoding="utf-8", errors="ignore") as f:
                    doc.raw_text = f.read()

            if not doc.raw_text:
                doc.status = "FAILED"
                await self.db.commit()
                return {"status": "FAILED", "error": ocr_res.error or "Text extraction failed"}
        else:
            doc.raw_text = ocr_res.text
            doc.ocr_engine = ocr_res.engine

        # Save OCR updates
        await self.db.commit()

        # 2. Document Classification
        classification_res = await self.classify_text(doc.raw_text, model)
        doc.classification = classification_res.get("classification")
        doc.classification_confidence = classification_res.get("confidence")

        # 3. Information Extraction
        schema = SCHEMAS.get(doc.classification, "{}")
        extracted = await self.extract_fields(doc.raw_text, schema, model)
        doc.extracted_data = extracted

        # 4. Field Validation
        val_results = DocumentValidator.validate_extracted_fields(doc.classification or "CUSTOM", extracted)
        doc.validation_results = val_results
        
        # Calculate Validation status
        invalid_count = sum(1 for v in val_results.values() if not v.get("valid", True))
        if not val_results:
            doc.validation_status = "UNVALIDATED"
        elif invalid_count == 0:
            doc.validation_status = "VALID"
        else:
            doc.validation_status = "INVALID"

        doc.status = "COMPLETED"
        await self.db.commit()

        await self.log_audit(
            user_id=doc.uploaded_by,
            doc_id=doc.id,
            action="CLASSIFY",
            details=f"Classified as {doc.classification} (conf: {doc.classification_confidence})"
        )

        return {
            "id": str(doc.id),
            "status": "COMPLETED",
            "classification": doc.classification,
            "confidence": doc.classification_confidence,
            "extracted_data": doc.extracted_data,
            "validation_status": doc.validation_status,
            "validation_results": doc.validation_results,
        }

    async def analyze_document_compliance(
        self,
        doc_uuid: uuid.UUID,
        model: Optional[str] = None,
    ) -> dict[str, Any]:
        """Compute Summary, Key Highlights, Risk, and AI recommendations."""
        from app.models.ai_document_analysis import AnalyzedDocument

        stmt = select(AnalyzedDocument).where(AnalyzedDocument.id == doc_uuid)
        res = await self.db.execute(stmt)
        doc = res.scalar_one_or_none()
        if not doc or not doc.raw_text:
            raise ValueError("Document not processed yet or empty raw text.")

        prompt = PromptLibrary.document_analysis_user(doc.raw_text)
        response = await self.llm.complete(
            prompt=prompt,
            system=PromptLibrary.DOCUMENT_ANALYSIS_SYSTEM,
            model=model,
            json_mode=True,
            temperature=0.2
        )
        data = ResponseParser.extract_json_object(response)

        # Write analysis to DB
        doc.summary_executive = data.get("summary_executive")
        doc.summary_detailed = data.get("summary_detailed")
        doc.key_highlights = data.get("key_highlights")
        doc.missing_info = data.get("missing_info")
        doc.compliance_report = data.get("compliance_report")
        doc.risk_analysis = data.get("risk_analysis")
        doc.ai_recommendations = data.get("ai_recommendations")
        doc.health_score = data.get("health_score", 1.0)

        await self.db.commit()

        await self.log_audit(
            user_id=doc.uploaded_by,
            doc_id=doc.id,
            action="ANALYZE",
            details="Executed full compliance and risk analysis"
        )

        return data

    async def compare_documents(
        self,
        left_uuid: uuid.UUID,
        right_uuid: uuid.UUID,
        user_uuid: Optional[uuid.UUID] = None,
        model: Optional[str] = None,
    ) -> dict[str, Any]:
        """Compare Document A and Document B for changes, additions, and deletions."""
        from app.models.ai_document_analysis import AnalyzedDocument, DocumentComparisonRun

        # Retrieve documents
        res_a = await self.db.execute(select(AnalyzedDocument).where(AnalyzedDocument.id == left_uuid))
        doc_a = res_a.scalar_one_or_none()
        res_b = await self.db.execute(select(AnalyzedDocument).where(AnalyzedDocument.id == right_uuid))
        doc_b = res_b.scalar_one_or_none()

        if not doc_a or not doc_b or not doc_a.raw_text or not doc_b.raw_text:
            raise ValueError("Both documents must exist and have processed raw text.")

        prompt = PromptLibrary.document_comparison_user(doc_a.raw_text, doc_b.raw_text)
        response = await self.llm.complete(
            prompt=prompt,
            system=PromptLibrary.DOCUMENT_COMPARISON_SYSTEM,
            model=model,
            json_mode=True,
            temperature=0.15
        )
        data = ResponseParser.extract_json_object(response)

        # Log comparison run to DB
        run = DocumentComparisonRun(
            source_document_id=left_uuid,
            target_document_id=right_uuid,
            similarity_score=data.get("similarity_score", 0.0),
            differences=data.get("differences"),
            missing_info=data.get("missing_info"),
            changed_fields=data.get("changed_fields"),
            fraud_signals=data.get("fraud_signals"),
            compared_by=user_uuid,
        )
        self.db.add(run)
        await self.db.commit()

        await self.log_audit(
            user_id=user_uuid,
            doc_id=left_uuid,
            action="COMPARE",
            details=f"Compared with document {right_uuid}. Similarity: {run.similarity_score:.0%}"
        )

        return {
            "comparison_run_id": str(run.id),
            "similarity_score": run.similarity_score,
            "differences": run.differences,
            "changed_fields": run.changed_fields,
            "missing_info": run.missing_info,
            "fraud_signals": run.fraud_signals,
        }

    # ------------------------------------------------------------------
    # LLM Wrapper Helpers
    # ------------------------------------------------------------------

    async def classify_text(self, text: str, model: Optional[str] = None) -> dict[str, Any]:
        """Classify raw text into a standard document category using LLM."""
        prompt = PromptLibrary.document_classification_user(text)
        response = await self.llm.complete(
            prompt=prompt,
            system=PromptLibrary.DOCUMENT_CLASSIFICATION_SYSTEM,
            model=model,
            json_mode=True,
            temperature=0.1
        )
        data = ResponseParser.extract_json_object(response)
        return {
            "classification": data.get("classification", "CUSTOM_DOCUMENT").upper(),
            "confidence": float(data.get("confidence", 0.5)),
        }

    async def extract_fields(self, text: str, schema_json: str, model: Optional[str] = None) -> dict[str, Any]:
        """Extract structured JSON fields from raw text based on specified schema."""
        prompt = PromptLibrary.document_extraction_user(text, schema_json)
        response = await self.llm.complete(
            prompt=prompt,
            system=PromptLibrary.DOCUMENT_EXTRACTION_SYSTEM,
            model=model,
            json_mode=True,
            temperature=0.1
        )
        return ResponseParser.extract_json_object(response)

    # ------------------------------------------------------------------
    # Fallback Parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_pdf_text_fallback(path: str) -> str:
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""

    @staticmethod
    def _extract_docx_text_fallback(path: str) -> str:
        try:
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""

    # ------------------------------------------------------------------
    # Audit Logs
    # ------------------------------------------------------------------

    async def log_audit(
        self,
        user_id: Optional[uuid.UUID],
        doc_id: Optional[uuid.UUID],
        action: str,
        details: str,
    ) -> None:
        """Create structured analysis audit entry."""
        from app.models.ai_document_analysis import AnalysisAuditLog
        log = AnalysisAuditLog(
            user_id=user_id,
            document_id=doc_id,
            action=action,
            details=details,
        )
        self.db.add(log)
        await self.db.commit()
