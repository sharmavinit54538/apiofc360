"""Service for extracting raw text from PDF, DOCX, DOC, TXT, and Image documents locally."""

from __future__ import annotations

import logging
import os
import re

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from PDF file with image OCR fallback for scanned pages."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text_list = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                text_list.append(t.strip())
            elif hasattr(page, "images") and page.images:
                for img in page.images:
                    img_text = _extract_image_bytes(img.data)
                    if img_text.strip():
                        text_list.append(img_text.strip())

        combined = "\n\n".join(text_list).strip()
        if combined:
            return combined
    except Exception as exc:
        logger.warning("extract_text_from_pdf failed | path=%s | exc=%s", file_path, exc)

    # Fallback string recovery
    try:
        with open(file_path, "rb") as f:
            content = f.read().decode("latin-1", errors="ignore")
        readable = re.findall(r"[A-Za-z0-9\s.,@+\-(){}:;/]{4,}", content)
        return "\n".join(readable).strip()
    except Exception:
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text, headings, and tables from DOCX document."""
    try:
        from docx import Document
        doc = Document(file_path)
        content_blocks = []

        for para in doc.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name or "title" in style_name:
                content_blocks.append(f"## {t}")
            else:
                content_blocks.append(t)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                unique_cells = []
                for c in cells:
                    if not unique_cells or unique_cells[-1] != c:
                        unique_cells.append(c)
                if unique_cells:
                    content_blocks.append(" | ".join(unique_cells))

        if content_blocks:
            return "\n".join(content_blocks)
    except Exception as exc:
        logger.debug("extract_text_from_docx failed | path=%s | exc=%s", file_path, exc)

    # Fallback zip XML reader
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(file_path) as z:
            xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            texts = [node.text for node in tree.iter() if node.text]
            return "\n".join(texts)
    except Exception:
        pass

    return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract raw text from plain text file using safe encoding detection."""
    encodings = ["utf-8", "utf-8-sig", "utf-16", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                content = f.read()
                if content and len(content.strip()) > 0:
                    return content.strip()
        except (UnicodeDecodeError, UnicodeError, OSError):
            continue
    try:
        with open(file_path, "rb") as f:
            return f.read().decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def extract_text_from_image(file_path: str) -> str:
    """Extract text from image files using OCR / PIL."""
    filename = os.path.basename(file_path)
    logger.info("OCR Image Extraction | path=%s", filename)
    try:
        from app.ocr.engine_selector import get_ocr_selector
        selector = get_ocr_selector()
        res = selector.extract_text_from_file(file_path)
        if res and res.success and len(res.text.strip()) > 10:
            return res.text.strip()
    except Exception as exc:
        logger.debug("Engine selector OCR failed: %s", exc)

    try:
        import pytesseract
        from PIL import Image
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as exc:
        logger.warning("OCR image extraction failed | path=%s | error=%s", filename, exc)
        return ""


def _extract_image_bytes(image_bytes: bytes) -> str:
    """Extract text from raw image bytes using OCR engine."""
    try:
        from app.ocr.engine_selector import get_ocr_selector
        selector = get_ocr_selector()
        res = selector.extract_text_from_image_bytes(image_bytes)
        if res and res.success:
            return res.text.strip()
    except Exception:
        pass
    return ""


def extract_document_text(file_path: str) -> str:
    """Detect file type and extract text cleanly."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in {".docx", ".doc"}:
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        return extract_text_from_image(file_path)
    return ""
