"""Resume Parser AI Agent.

Full production pipeline:
1. Accept PDF / DOCX / DOC / JPG / PNG / JPEG / TIFF
2. Route to OCR engine (images) or document parser (PDF/DOCX)
3. Send extracted text to Ollama LLM for structured extraction
4. Extract 25+ fields including skills taxonomy, experience timeline,
   certifications, projects, publications, awards, achievements
5. Return validated structured JSON

All fields stored as JSON in the database for full-text and semantic search.
"""

from __future__ import annotations

import asyncio
import logging
import os
from typing import Any

from app.llm.client import LLMClient, get_llm_client
from app.llm.prompts import PromptLibrary
from app.llm.response_parser import ResponseParser
from app.ocr.engine_selector import OCREngineSelector, get_ocr_selector

logger = logging.getLogger(__name__)

# Supported file extensions
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif"}
DOC_EXTS = {".pdf", ".doc", ".docx"}


class ParsedResume:
    """Value object holding all extracted resume fields."""

    def __init__(self, data: dict[str, Any], raw_text: str, engine_used: str) -> None:
        self.raw_data = data
        self.raw_text = raw_text
        self.engine_used = engine_used  # 'paddle'|'easyocr'|'tesseract'|'pypdf'|'docx'

        # Core identity
        self.name: str | None = data.get("name")
        self.email: str | None = data.get("email")
        self.mobile: str | None = data.get("mobile")
        self.address: str | None = data.get("address")
        self.linkedin_url: str | None = data.get("linkedin_url")
        self.github_url: str | None = data.get("github_url")
        self.portfolio_url: str | None = data.get("portfolio_url")

        # Professional context
        self.current_company: str | None = data.get("current_company")
        self.current_designation: str | None = data.get("current_designation")
        self.current_salary: float | None = self._to_float(data.get("current_salary"))
        self.expected_salary: float | None = self._to_float(data.get("expected_salary"))
        self.notice_period: str | None = data.get("notice_period")
        self.years_experience: float = self._to_float(data.get("years_experience")) or 0.0

        # Languages spoken
        self.languages: list[str] = ResponseParser.get_list(data, "languages")

        # Summary
        self.summary: str | None = data.get("summary")

        # Skills (taxonomized)
        skills_raw = data.get("skills") or {}
        if isinstance(skills_raw, list):
            # LLM returned flat list instead of object
            self.skills: dict[str, list[str]] = {"other": skills_raw}
        else:
            self.skills = {
                "programming_languages": self._coerce_list(skills_raw.get("programming_languages")),
                "frameworks": self._coerce_list(skills_raw.get("frameworks")),
                "databases": self._coerce_list(skills_raw.get("databases")),
                "cloud": self._coerce_list(skills_raw.get("cloud")),
                "tools": self._coerce_list(skills_raw.get("tools")),
                "soft_skills": self._coerce_list(skills_raw.get("soft_skills")),
                "domain_expertise": self._coerce_list(skills_raw.get("domain_expertise")),
                "other": self._coerce_list(skills_raw.get("other")),
            }

        # Experience timeline
        self.experience: list[dict] = ResponseParser.get_list(data, "experience")
        self.previous_companies: list[str] = ResponseParser.get_list(data, "previous_companies")

        # Education
        self.education: list[dict] = ResponseParser.get_list(data, "education")

        # Certifications, projects, publications, awards
        self.certifications: list[dict] = ResponseParser.get_list(data, "certifications")
        self.projects: list[dict] = ResponseParser.get_list(data, "projects")
        self.publications: list[str] = ResponseParser.get_list(data, "publications")
        self.awards: list[str] = ResponseParser.get_list(data, "awards")
        self.achievements: list[str] = ResponseParser.get_list(data, "achievements")

    @property
    def all_skills_flat(self) -> list[str]:
        """Return deduplicated flat list of all skills."""
        seen: set[str] = set()
        result: list[str] = []
        for skill_list in self.skills.values():
            for skill in skill_list:
                normalized = str(skill).strip().lower()
                if normalized and normalized not in seen:
                    seen.add(normalized)
                    result.append(str(skill).strip())
        return result

    def to_dict(self) -> dict[str, Any]:
        """Serialize to dict for DB storage."""
        return {
            "name": self.name,
            "email": self.email,
            "mobile": self.mobile,
            "address": self.address,
            "linkedin_url": self.linkedin_url,
            "github_url": self.github_url,
            "portfolio_url": self.portfolio_url,
            "current_company": self.current_company,
            "current_designation": self.current_designation,
            "current_salary": self.current_salary,
            "expected_salary": self.expected_salary,
            "notice_period": self.notice_period,
            "years_experience": self.years_experience,
            "languages": self.languages,
            "summary": self.summary,
            "skills": self.skills,
            "experience": self.experience,
            "previous_companies": self.previous_companies,
            "education": self.education,
            "certifications": self.certifications,
            "projects": self.projects,
            "publications": self.publications,
            "awards": self.awards,
            "achievements": self.achievements,
            "engine_used": self.engine_used,
        }

    @staticmethod
    def _to_float(val: Any) -> float | None:
        if val is None:
            return None
        try:
            return float(str(val).replace(",", "").replace("$", "").replace("₹", "").strip())
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _coerce_list(val: Any) -> list[str]:
        if isinstance(val, list):
            return [str(v).strip() for v in val if v]
        if isinstance(val, str) and val:
            return [v.strip() for v in val.split(",") if v.strip()]
        return []


class ResumeParserAgent:
    """Full AI resume parser supporting all document and image types.

    Orchestrates:
    - PDF/DOCX text extraction via pypdf/python-docx
    - Image OCR via PaddleOCR / EasyOCR / Tesseract
    - LLM-based structured data extraction
    - Validation and normalization
    """

    def __init__(
        self,
        llm_client: LLMClient | None = None,
        ocr_selector: OCREngineSelector | None = None,
    ) -> None:
        self._llm = llm_client or get_llm_client()
        self._ocr = ocr_selector or get_ocr_selector()

    async def parse_file(self, file_path: str, model: str | None = None) -> ParsedResume:
        """Parse a resume file and return structured extracted data.

        Supports: PDF, DOCX, DOC, JPG, JPEG, PNG, TIFF
        """
        _, ext = os.path.splitext(file_path.lower())

        # Step 1: Extract raw text
        raw_text, engine_used = await self._extract_raw_text(file_path, ext)

        if not raw_text or len(raw_text.strip()) < 30:
            logger.warning("Resume text too short (%d chars) for %s", len(raw_text), file_path)
            return ParsedResume({}, raw_text, engine_used)

        # Step 2: LLM extraction
        parsed_data = await self._llm_extract(raw_text, model=model)

        return ParsedResume(parsed_data, raw_text, engine_used)

    async def parse_text(self, text: str, model: str | None = None) -> ParsedResume:
        """Parse pre-extracted resume text directly (skip file I/O)."""
        if not text or len(text.strip()) < 20:
            return ParsedResume({}, text, "direct")
        parsed_data = await self._llm_extract(text, model=model)
        return ParsedResume(parsed_data, text, "direct")

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    async def _extract_raw_text(self, file_path: str, ext: str) -> tuple[str, str]:
        """Extract raw text from file based on extension."""
        if ext in IMAGE_EXTS:
            # OCR path (runs in thread pool to avoid blocking event loop)
            result = await asyncio.get_event_loop().run_in_executor(
                None, self._ocr.extract_text_from_file, file_path
            )
            return result.text, result.engine_used

        elif ext == ".pdf":
            text = await asyncio.get_event_loop().run_in_executor(
                None, self._extract_pdf, file_path
            )
            return text, "pypdf"

        elif ext in {".docx", ".doc"}:
            text = await asyncio.get_event_loop().run_in_executor(
                None, self._extract_docx, file_path
            )
            return text, "python-docx"

        else:
            logger.warning("Unsupported file extension: %s", ext)
            return "", "none"

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from PDF using pypdf."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            return "\n".join(pages)
        except Exception as exc:
            logger.error("PDF extraction failed for %s: %s", file_path, exc)
            return ""

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as exc:
            logger.error("DOCX extraction failed for %s: %s", file_path, exc)
            return ""

    async def _llm_extract(self, raw_text: str, model: str | None = None) -> dict[str, Any]:
        """Send raw text to Ollama LLM for structured extraction."""
        # Sanitize input against prompt injection
        safe_text = ResponseParser.sanitize_user_input(raw_text, max_length=8000)

        prompt = PromptLibrary.resume_parser_user(safe_text)
        system = PromptLibrary.RESUME_PARSER_SYSTEM

        response = await self._llm.complete(
            prompt=prompt,
            system=system,
            model=model,
            json_mode=True,
            num_predict=3000,
            temperature=0.1,  # Low temp for extraction tasks
        )

        if not response:
            logger.warning("LLM returned empty response for resume parsing")
            return {}

        parsed = ResponseParser.extract_json_object(response)
        if not parsed:
            logger.warning("Could not extract JSON from LLM resume response")
        return parsed
